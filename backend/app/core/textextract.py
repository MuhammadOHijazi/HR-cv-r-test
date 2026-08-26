"""Turn a raw CV file into plain text plus a source-quality score.

PDFs go through PyMuPDF.  A PDF whose text layer is empty or near-empty is
treated as *scanned*: its pages are rasterised and handed to the Gemini vision
fallback through the one Gemini gateway.  DOCX goes through python-docx.

``source_quality`` in [0,1] feeds the confidence assembly: a clean digital PDF
scores high, an OCR'd scan scores low, and low quality is a routing flag.
"""

from __future__ import annotations

import logging
import re
import unicodedata
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Protocol

logger = logging.getLogger(__name__)

PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# A page with fewer than this many extractable characters is treated as image-only.
SCANNED_CHARS_PER_PAGE = 40
MAX_VISION_PAGES = 10


class VisionFallback(Protocol):
    """The seam to the Gemini vision path (implemented by :mod:`ingestion`)."""

    def transcribe(self, images: list[bytes]) -> str: ...


@dataclass
class ExtractedText:
    text: str
    page_count: int
    is_scanned: bool
    source_quality: float
    method: str
    warnings: list[str] = field(default_factory=list)

    def as_dict(self) -> dict[str, Any]:
        return {
            "page_count": self.page_count,
            "is_scanned": self.is_scanned,
            "source_quality": round(self.source_quality, 4),
            "method": self.method,
            "warnings": self.warnings,
            "chars": len(self.text),
        }


class UnsupportedFormat(ValueError):
    """The file is not a format this pipeline reads."""


# Arabic presentation forms (U+FB50-U+FEFF) plus the zero-width joiners that
# PDF producers scatter through shaped Arabic runs.
_PRESENTATION_FORMS = re.compile(r"[ﭐ-﷿ﹰ-ﻼ]")
_ZERO_WIDTH = re.compile(r"[​-‏‪-‮⁦-⁩﻿]")


def normalise_extracted(text: str) -> str:
    """Clean up text as it comes out of a PDF or DOCX.

    Many PDF producers store Arabic as *presentation forms* (the shaped glyph
    codepoints) rather than the base letters.  NFKC folds those back onto the
    base Arabic block, which is what the taxonomy, the regex layer and the
    evidence verifier all expect.  Bidi control characters and non-breaking
    spaces are stripped for the same reason.
    """
    if not text:
        return ""
    out = _ZERO_WIDTH.sub("", text)
    if _PRESENTATION_FORMS.search(out):
        out = unicodedata.normalize("NFKC", out)
    out = out.replace("\xa0", " ")
    return out


def has_presentation_forms(text: str) -> bool:
    return bool(_PRESENTATION_FORMS.search(text or ""))


def detect_mime(filename: str, declared: str | None = None) -> str:
    if declared in {PDF_MIME, DOCX_MIME}:
        return declared
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return PDF_MIME
    if suffix == ".docx":
        return DOCX_MIME
    raise UnsupportedFormat(f"unsupported CV format: {filename!r} ({declared})")


def extract(
    data: bytes,
    filename: str,
    *,
    declared_mime: str | None = None,
    vision: VisionFallback | None = None,
) -> ExtractedText:
    """Extract text from raw file bytes."""
    mime = detect_mime(filename, declared_mime)
    if mime == PDF_MIME:
        return _extract_pdf(data, vision=vision)
    return _extract_docx(data)


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


def _open_pdf(data: bytes):
    import pymupdf  # PyMuPDF >= 1.25 exposes itself as `pymupdf`

    return pymupdf.open(stream=data, filetype="pdf")


def _extract_pdf(data: bytes, *, vision: VisionFallback | None) -> ExtractedText:
    warnings: list[str] = []
    try:
        doc = _open_pdf(data)
    except Exception as exc:
        raise UnsupportedFormat(f"could not open PDF: {exc}") from exc

    with doc:
        page_count = doc.page_count
        pages = [doc.load_page(i).get_text("text") or "" for i in range(page_count)]
        text = normalise_extracted("\n".join(pages)).strip()
        chars_per_page = len(text) / page_count if page_count else 0.0
        scanned = page_count > 0 and chars_per_page < SCANNED_CHARS_PER_PAGE

        if not scanned:
            return ExtractedText(
                text=text,
                page_count=page_count,
                is_scanned=False,
                source_quality=_quality_for_digital(text, page_count),
                method="pymupdf",
                warnings=warnings,
            )

        warnings.append("no_text_layer")
        if vision is None:
            warnings.append("vision_fallback_unavailable")
            return ExtractedText(
                text=text,
                page_count=page_count,
                is_scanned=True,
                source_quality=0.1,
                method="pymupdf-empty",
                warnings=warnings,
            )
        images = _render_pages(doc, min(page_count, MAX_VISION_PAGES))
        if page_count > MAX_VISION_PAGES:
            warnings.append(f"vision_truncated_to_{MAX_VISION_PAGES}_pages")

    try:
        vision_text = vision.transcribe(images)
    except Exception as exc:
        logger.warning("vision fallback failed: %s", exc)
        warnings.append("vision_fallback_failed")
        return ExtractedText("", page_count, True, 0.1, "vision-failed", warnings)

    return ExtractedText(
        text=normalise_extracted(vision_text).strip(),
        page_count=page_count,
        is_scanned=True,
        source_quality=_quality_for_vision(vision_text, page_count),
        method="gemini-vision",
        warnings=warnings,
    )


def _render_pages(doc, count: int) -> list[bytes]:
    images: list[bytes] = []
    for i in range(count):
        pix = doc.load_page(i).get_pixmap(dpi=150)
        images.append(pix.tobytes("png"))
    return images


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def _extract_docx(data: bytes) -> ExtractedText:
    import io

    from docx import Document

    try:
        document = Document(io.BytesIO(data))
    except Exception as exc:
        raise UnsupportedFormat(f"could not open DOCX: {exc}") from exc

    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = normalise_extracted("\n".join(p for p in parts if p is not None)).strip()
    pages = max(1, len(text) // 2800 + (1 if len(text) % 2800 else 0))
    return ExtractedText(
        text=text,
        page_count=pages,
        is_scanned=False,
        source_quality=_quality_for_digital(text, pages),
        method="python-docx",
    )


# ---------------------------------------------------------------------------
# Quality scoring
# ---------------------------------------------------------------------------

_WORDLIKE = re.compile(r"[A-Za-z؀-ۿ]{2,}")
_REPLACEMENT = re.compile("[\ufffd\x00-\x08\x0b\x0c\x0e-\x1f]")


def _quality_for_digital(text: str, page_count: int) -> float:
    """Quality of a native text layer: length, word density and cleanliness."""
    if not text.strip():
        return 0.05
    words = _WORDLIKE.findall(text)
    density = len(words) / max(len(text), 1)
    length_score = min(1.0, len(text) / (600.0 * max(page_count, 1)))
    density_score = min(1.0, density / 0.13)
    noise = len(_REPLACEMENT.findall(text)) / max(len(text), 1)
    noise_penalty = min(0.4, noise * 20)
    return round(max(0.05, 0.45 * length_score + 0.55 * density_score - noise_penalty), 4)


def _quality_for_vision(text: str, page_count: int) -> float:
    """OCR output is never trusted as much as a native text layer."""
    if not text.strip():
        return 0.1
    base = _quality_for_digital(text, page_count)
    return round(min(0.6, max(0.15, base * 0.6)), 4)
